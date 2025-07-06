from docx import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from transformers import pipeline
import nltk
nltk.download('punkt_tab')

#CONFIG 
INSTRUCTION_DOCX = "instructions.docx"
CONTENT_DOCX = "conents.docx"
CHUNK_SIZE = 5  # Number of sentences per chunk (tunable)
EMBEDDING_MODEL = 'all-MiniLM-L6-v2'
GENERATION_MODEL = 'tiiuae/falcon-7b-instruct'  # Can be replaced with Ollama/llama2 if needed



# READ DOCX CONTENT
def extract_text_from_docx(path):
    doc = Document(path)
    return '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

instruction_text = extract_text_from_docx(INSTRUCTION_DOCX)
content_text = extract_text_from_docx(CONTENT_DOCX)


#CHUNK THE TEXT
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

def chunk_text(text, chunk_size=CHUNK_SIZE):
    sentences = sent_tokenize(text)
    return [' '.join(sentences[i:i+chunk_size]) for i in range(0, len(sentences), chunk_size)]

instruction_chunks = chunk_text(instruction_text, 1)  # One instruction per chunk
content_chunks = chunk_text(content_text, CHUNK_SIZE)


# EMBEDDINGS
model = SentenceTransformer(EMBEDDING_MODEL)
instruction_vectors = model.encode(instruction_chunks)
content_vectors = model.encode(content_chunks)


# MATCH INSTRUCTIONS 
def find_best_match(instruction_vector, content_vectors):
    sims = cosine_similarity([instruction_vector], content_vectors)[0]
    return sims.argmax(), sims.max()

instruction_to_content_map = []
for idx, instr_vec in enumerate(instruction_vectors):
    best_match_idx, sim_score = find_best_match(instr_vec, content_vectors)
    instruction_to_content_map.append((idx, best_match_idx, sim_score))


#MODIFY TEXT

# Replace this with Ollama2 pipeline if you're using Ollama locally
#generator = pipeline("text2text-generation", model=GENERATION_MODEL)

import requests

def ollama_generate(prompt):
    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": "llama2",  # Make sure this matches your pulled model
            "prompt": prompt,
            "stream": False
        }
    )

    print("Ollama raw response:", response.json())  # Add this line

    return response.json()["response"]


modified_chunks = content_chunks.copy()

for instr_idx, cont_idx, score in instruction_to_content_map:
    instruction = instruction_chunks[instr_idx]
    original = content_chunks[cont_idx]
    
    prompt = f"""Apply the following instruction to the text:\n\nInstruction: {instruction}\n\nText: {original}\n\nModified Text:"""
    
    #response = generator(prompt, max_new_tokens=200)[0]['generated_text']
    response = ollama_generate(prompt)

    
    # Replace original with modified version
    modified_chunks[cont_idx] = response


# RECONSTRUCT OUTPUT 
from docx import Document

output_doc = Document()
output_doc.add_heading('Modified Content', level=1)

for chunk in modified_chunks:
    output_doc.add_paragraph(chunk)

output_doc.save("modified_output.docx")
print("Modified document saved as 'modified_output.docx'")
